"""
rag/document_loader.py
----------------------
Load PDF, DOCX, TXT files thành danh sách Document của LangChain.
Tối ưu hóa chunking sử dụng Semantic Separators ranh giới ngữ nghĩa.

Hỗ trợ 3 nguồn:
  1. load(file_path)        — file trên disk (PDF/DOCX/TXT)
  2. load_from_text(text)   — text thuần (Google Docs/Slides/Sheets export)
  3. load_from_bytes(bytes) — binary content (PDF/DOCX tải từ Drive)
"""
import os
import tempfile
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from core.logger import logger

# Chunk size có thể lớn hơn do đã có bộ chia cắt tốt hơn
CHUNK_SIZE = 1000       # ký tự per chunk
CHUNK_OVERLAP = 200     # overlap để không mất context ở ranh giới


class DocumentLoader:
    """Load và split tài liệu thành chunks theo ngữ nghĩa của câu/đoạn."""

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
    ):
        # Dùng regex để cắt đúng dấu chấm nhưng không mất dấu chấm
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n",
                "\n",
                r"(?<=\. )",  # Cắt theo dấu chấm câu tiếng Việt/Anh
                " ",
                ""
            ],
            is_separator_regex=True
        )

    # ── Source 1: File trên disk ─────────────────────────────────────────

    def load(self, file_path: str) -> List[Document]:
        """Load file → split → trả về list Document."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

        ext = path.suffix.lower()
        logger.info(f"DocumentLoader: loading {path.name} ({ext})")

        if ext == ".pdf":
            raw_docs = self._load_pdf(path)
        elif ext == ".docx":
            raw_docs = self._load_docx(path)
        elif ext == ".txt":
            raw_docs = self._load_txt(path)
        else:
            raise ValueError(f"Định dạng không hỗ trợ: {ext}. Chỉ nhận PDF, DOCX, TXT.")

        chunks = self._splitter.split_documents(raw_docs)
        logger.info(f"DocumentLoader: {len(chunks)} chunks từ {path.name}")
        return chunks

    # ── Source 2: Text thuần (Google Docs/Slides export) ─────────────────

    def load_from_text(self, text: str, metadata: dict) -> List[Document]:
        """
        Load từ text string (Google Drive export).
        Dùng cho: Google Docs, Google Slides, Google Sheets (CSV), TXT từ Drive.

        Args:
            text: Nội dung text đã extract.
            metadata: Dict chứa source, drive_file_id, drive_modified_time, v.v.
        Returns:
            Danh sách Document chunks.
        """
        if not text or not text.strip():
            logger.warning(f"DocumentLoader.load_from_text: text rỗng cho {metadata.get('source', '?')}")
            return []

        doc = Document(page_content=text, metadata=metadata)
        chunks = self._splitter.split_documents([doc])
        # Đảm bảo metadata được copy sang từng chunk
        for chunk in chunks:
            chunk.metadata.update(metadata)

        logger.info(f"DocumentLoader.load_from_text: {len(chunks)} chunks từ '{metadata.get('source', '?')}'")
        return chunks

    # ── Source 3: Binary content (PDF/DOCX tải từ Drive) ─────────────────

    def load_from_bytes(self, content: bytes, ext: str, metadata: dict) -> List[Document]:
        """
        Load từ binary content (file download từ Google Drive).
        Dùng cho: PDF, DOCX tải từ Drive.

        Args:
            content: Nội dung binary của file.
            ext: Extension file (".pdf", ".docx", ".txt").
            metadata: Dict chứa source, drive_file_id, v.v.
        Returns:
            Danh sách Document chunks.
        """
        if not content:
            logger.warning(f"DocumentLoader.load_from_bytes: content rỗng cho {metadata.get('source', '?')}")
            return []

        ext = ext.lower()
        logger.info(f"DocumentLoader.load_from_bytes: xử lý {ext} ({len(content)} bytes)")

        # Ghi ra file tạm để dùng lại các parser hiện có
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(content)
            tmp_path = Path(tmp.name)

        try:
            if ext == ".pdf":
                raw_docs = self._load_pdf(tmp_path)
            elif ext == ".docx":
                raw_docs = self._load_docx(tmp_path)
            elif ext in (".txt", ".csv"):
                raw_docs = self._load_txt(tmp_path)
            else:
                raise ValueError(f"Extension không hỗ trợ: {ext}")

            # Ghi đè metadata với thông tin Drive
            for doc in raw_docs:
                doc.metadata.update(metadata)

            chunks = self._splitter.split_documents(raw_docs)
            for chunk in chunks:
                chunk.metadata.update(metadata)

            logger.info(f"DocumentLoader.load_from_bytes: {len(chunks)} chunks")
            return chunks
        finally:
            tmp_path.unlink(missing_ok=True)

    # ── Private parsers ──────────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> List[Document]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "page": i + 1, "file_path": str(path)},
                ))
        return docs

    def _load_docx(self, path: Path) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(
            page_content=text,
            metadata={"source": path.name, "file_path": str(path)},
        )]

    def _load_txt(self, path: Path) -> List[Document]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [Document(
            page_content=text,
            metadata={"source": path.name, "file_path": str(path)},
        )]


# ── Module-level helper (backward compatibility) ─────────────────────────────

def load_document(file_path: str) -> List[Document]:
    """Backward-compatible helper. Dùng DocumentLoader.load() ngầm định."""
    return DocumentLoader().load(file_path)
