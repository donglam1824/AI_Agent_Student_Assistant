"""
rag/document_loader.py
----------------------
Load PDF, DOCX, TXT files thành danh sách Document của LangChain.
Tối ưu hóa chunking sử dụng Semantic Separators ranh giới ngữ nghĩa.
"""
import os
from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from core.logger import logger

# Chunk size có thể lớn hơn do đã có bộ chia cắt tốt hơn
CHUNK_SIZE = 1000       # ký tự per chunk
CHUNK_OVERLAP = 200     # overlap để không mất context ở ranh giới


class DocumentLoader:
    """Load và split tài liệu thành chunks theo ngữ nghĩa của câu/đoạn."""

    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP,
    ):
        # Dùng regex để cắt đúng dấu chấm nhưng không mất dấu chấm
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=[
                "\n\n",
                "\n",
                r"(?<=\. )",  # Cắt theo dấu chấm câu tiếng Việt/Anh
                " ",
                ""
            ],
            is_separator_regex=True
        )

    def load(self, file_path: str) -> List[Document]:
        """Load file → split → trả về list Document."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Không tìm thấy file: {file_path}")

        ext = path.suffix.lower()
        logger.info(f"DocumentLoader: loading {path.name} ({ext})")

        if ext == ".pdf":
            raw_docs = self._load_pdf(path)
        elif ext == ".docx":
            raw_docs = self._load_docx(path)
        elif ext == ".txt":
            raw_docs = self._load_txt(path)
        else:
            raise ValueError(f"Định dạng không hỗ trợ: {ext}. Chỉ nhận PDF, DOCX, TXT.")

        chunks = self._splitter.split_documents(raw_docs)
        logger.info(f"DocumentLoader: {len(chunks)} chunks từ {path.name}")
        return chunks

    def _load_pdf(self, path: Path) -> List[Document]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        docs = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": path.name, "page": i + 1, "file_path": str(path)},
                ))
        return docs

    def _load_docx(self, path: Path) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(
            page_content=text,
            metadata={"source": path.name, "file_path": str(path)},
        )]

    def _load_txt(self, path: Path) -> List[Document]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return [Document(
            page_content=text,
            metadata={"source": path.name, "file_path": str(path)},
        )]
